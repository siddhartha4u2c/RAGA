"""
Multilingual Indian Language Chatbot - Session-only memory, text + voice I/O.
Deployable on Render via GitHub.
"""
import base64
import json
import os
import uuid
from datetime import datetime, timezone, timedelta
from pathlib import Path
from urllib.request import Request, urlopen
from urllib.parse import quote, urlparse
from urllib.error import URLError, HTTPError

# Load .env from project root (for local runs)
if Path(__file__).resolve().parent.joinpath(".env").exists():
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parent / ".env")

from flask import Flask, request, jsonify, session, render_template

# Optional: transliteration for Indian languages (Latin <-> Indic script)
import re
SCRIPT_MAP = {}
INDIC_BLOCKS = {
    "hi": re.compile(r"[\u0900-\u097F]"),   # Devanagari
    "mr": re.compile(r"[\u0900-\u097F]"),
    "ta": re.compile(r"[\u0B80-\u0BFF]"),   # Tamil
    "te": re.compile(r"[\u0C00-\u0C7F]"),   # Telugu
    "as": re.compile(r"[\u0980-\u09FF]"),   # Assamese (same block as Bengali)
    "bn": re.compile(r"[\u0980-\u09FF]"),   # Bengali
    "gu": re.compile(r"[\u0A80-\u0AFF]"),   # Gujarati
    "kn": re.compile(r"[\u0C80-\u0CFF]"),   # Kannada
    "ml": re.compile(r"[\u0D00-\u0D7F]"),   # Malayalam
    "pa": re.compile(r"[\u0A00-\u0A7F]"),   # Gurmukhi
}
try:
    from indic_transliteration import sanscript
    from indic_transliteration.sanscript import transliterate
    HAS_TRANSLITERATION = True
    SCRIPT_MAP = {
        "hi": sanscript.DEVANAGARI,
        "ta": sanscript.TAMIL,
        "te": sanscript.TELUGU,
        "as": sanscript.BENGALI,   # Assamese uses Bengali script
        "bn": sanscript.BENGALI,
        "mr": sanscript.DEVANAGARI,
        "gu": sanscript.GUJARATI,
        "kn": sanscript.KANNADA,
        "ml": sanscript.MALAYALAM,
        "pa": sanscript.GURMUKHI,
        "en": None,
    }
except ImportError:
    HAS_TRANSLITERATION = False

# In-memory session store: session_id -> list of {role, content}
# Memory is local: cleared when user clicks "New chat" or changes language (api_clear creates a new session).
SESSION_MEMORY: dict[str, list[dict]] = {}
# Uploaded files per session: session_id -> list of {type, content, filename, mime?}
SESSION_UPLOADS: dict[str, list[dict]] = {}

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "dev-secret-change-in-production")


def get_or_create_session_id():
    if "session_id" not in session:
        session["session_id"] = str(uuid.uuid4())
        SESSION_MEMORY[session["session_id"]] = []
    return session["session_id"]


def get_messages(session_id: str):
    return SESSION_MEMORY.get(session_id, [])


def add_message(session_id: str, role: str, content: str):
    if session_id not in SESSION_MEMORY:
        SESSION_MEMORY[session_id] = []
    SESSION_MEMORY[session_id].append({"role": role, "content": content})
    # Keep only last N turns to limit memory (e.g. last 20 messages)
    max_messages = 50
    if len(SESSION_MEMORY[session_id]) > max_messages:
        SESSION_MEMORY[session_id] = SESSION_MEMORY[session_id][-max_messages:]


def _has_indic_script(text: str, lang_code: str) -> bool:
    """True if text contains characters from the Indic script for this language."""
    pat = INDIC_BLOCKS.get(lang_code)
    return bool(pat and pat.search(text))


# Any Indic script (Devanagari, Tamil, Telugu, Bengali, etc.) - for enforcing English-only replies
INDIC_SCRIPT_PATTERN = re.compile(
    r"[\u0900-\u097F\u0980-\u09FF\u0B80-\u0BFF\u0C00-\u0C7F\u0A80-\u0AFF\u0C80-\u0CFF\u0D00-\u0D7F\u0A00-\u0A7F]"
)


def _reply_contains_indic(text: str) -> bool:
    """True if text contains Indic script characters (so reply is not English-only)."""
    return bool(text and INDIC_SCRIPT_PATTERN.search(text))


# Assamese uses ৰ (U+09F0) and ৱ (U+09F1); Bengali uses র (U+09B0). Do not mix.
BENGALI_RA = "\u09B0"
ASSAMESE_RA = "\u09F0"
ASSAMESE_VA = "\u09F1"


def _bengali_to_assamese(text: str) -> str:
    """Replace Bengali-specific letters with Assamese equivalents."""
    if not text:
        return text
    return text.replace(BENGALI_RA, ASSAMESE_RA)


def _assamese_to_bengali_for_roman(text: str) -> str:
    """Replace Assamese ৰ so library can convert to Roman (library expects Bengali)."""
    if not text:
        return text
    return text.replace(ASSAMESE_RA, BENGALI_RA)


def transliterate_to_indic(text: str, lang_code: str) -> str:
    """Convert Latin/English-like input to Indian script if supported."""
    if not text or not HAS_TRANSLITERATION or lang_code not in SCRIPT_MAP:
        return text
    target_script = SCRIPT_MAP.get(lang_code)
    if not target_script:
        return text
    try:
        out = transliterate(text, sanscript.ITRANS, target_script)
        if lang_code == "as":
            out = _bengali_to_assamese(out)
        return out
    except Exception:
        return text


def transliterate_to_roman(text: str, lang_code: str) -> str:
    """Convert Indic script to Roman (ITRANS) for the given language."""
    if not text or not HAS_TRANSLITERATION or lang_code not in SCRIPT_MAP:
        return text
    source_script = SCRIPT_MAP.get(lang_code)
    if not source_script:
        return text
    try:
        if lang_code == "as":
            text = _assamese_to_bengali_for_roman(text)
        return transliterate(text, source_script, sanscript.ITRANS)
    except Exception:
        return text


def extract_pdf_text(data: bytes) -> str:
    try:
        from pypdf import PdfReader
        from io import BytesIO
        reader = PdfReader(BytesIO(data))
        parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
        return "\n\n".join(parts) if parts else ""
    except Exception:
        return ""


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs) if doc.paragraphs else ""
    except Exception:
        return ""


# URL detection and web page content extraction for shared links
URL_PATTERN = re.compile(
    r"https?://[^\s<>\"']+",
    re.IGNORECASE,
)
MAX_LINKS_PER_MESSAGE = 3
LINK_FETCH_TIMEOUT = 15
LINK_CONTENT_MAX_CHARS = 40000

# Browser-like User-Agent and headers so sites (e.g. news) are less likely to block
LINK_FETCH_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Accept-Encoding": "gzip, deflate, br",
    "DNT": "1",
    "Connection": "keep-alive",
    "Upgrade-Insecure-Requests": "1",
}


def _fetch_url_html(url: str) -> str | None:
    """Fetch URL and return raw HTML or None. Tries trafilatura, then requests with session, then urllib."""
    url = (url or "").strip()
    if not url.startswith("http://") and not url.startswith("https://"):
        url = "https://" + url
    # 1) Trafilatura's fetch (designed for article/news pages)
    try:
        import trafilatura
        from trafilatura.downloads import fetch_url
        raw = fetch_url(url, config=None)
        if raw and len(raw.strip()) > 200:
            return raw
    except Exception:
        pass
    # 2) requests with Session (cookies + redirects)
    try:
        import requests
        session = requests.Session()
        session.headers.update(LINK_FETCH_HEADERS)
        r = session.get(url, timeout=LINK_FETCH_TIMEOUT, allow_redirects=True)
        r.raise_for_status()
        r.encoding = r.apparent_encoding or "utf-8"
        if r.text and len(r.text.strip()) > 200:
            return r.text
    except Exception:
        pass
    # 3) urllib fallback
    try:
        req = Request(url, headers=LINK_FETCH_HEADERS)
        with urlopen(req, timeout=LINK_FETCH_TIMEOUT) as resp:
            raw = resp.read()
            charset = resp.headers.get_content_charset() or "utf-8"
            return raw.decode(charset, errors="replace")
    except Exception:
        return None


def _extract_text_with_bs4(html: str) -> str:
    """Use BeautifulSoup to extract text, preferring article/main/content areas (good for news sites)."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript", "iframe"]):
            tag.decompose()
        # Prefer main content: article, main, common content selectors
        for selector in ("article", "main", "[role='main']", ".post", ".content", ".article-body", ".entry-content", ".story-content", "#content", ".page-content"):
            container = soup.select_one(selector)
            if container:
                text = container.get_text(separator="\n", strip=True)
                if text and len(text) > 150:
                    return text
        # Fallback: remove nav/footer and get all text
        for tag in soup(["nav", "footer", "header", "aside"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except Exception:
        return ""


def _extract_text_from_html(html: str) -> str:
    """Extract main text from HTML. Use trafilatura first, then BeautifulSoup (article-aware)."""
    text = ""
    try:
        import trafilatura
        text = trafilatura.extract(html, include_comments=False, include_tables=True) or ""
    except Exception:
        pass
    if not text or len(text.strip()) < 100:
        text = _extract_text_with_bs4(html)
    return (text or "").strip()


def fetch_link_content(url: str) -> str | None:
    """Fetch URL and return extracted plain text, or None on failure."""
    html = _fetch_url_html(url)
    if not html:
        return None
    text = _extract_text_from_html(html)
    if not text:
        return None
    return text[:LINK_CONTENT_MAX_CHARS]


def extract_urls_from_text(text: str) -> list[str]:
    """Return list of unique HTTP(S) URLs found in text."""
    seen = set()
    out = []
    for m in URL_PATTERN.finditer(text):
        url = m.group(0).rstrip(".,;:)")
        if url not in seen:
            seen.add(url)
            out.append(url)
    return out


def _domain_from_url(url: str) -> str:
    """Extract host (e.g. makemytrip.com) from URL for search context."""
    try:
        p = urlparse(url if url.startswith("http") else "https://" + url)
        host = (p.netloc or "").strip().lower()
        if host.startswith("www."):
            host = host[4:]
        return host if host else ""
    except Exception:
        return ""


def _linkedin_profile_id(url: str) -> str | None:
    """Extract LinkedIn profile username from URL (e.g. linkedin.com/in/sagarkamble01 -> sagarkamble01)."""
    if not url or "linkedin.com" not in url.lower() or "/in/" not in url.lower():
        return None
    try:
        url = url if url.startswith("http") else "https://" + url
        p = urlparse(url)
        path = (p.path or "").strip().rstrip("/")
        if "/in/" in path.lower():
            segment = path.lower().split("/in/")[-1].split("/")[0].split("?")[0].strip()
            return segment if segment and len(segment) < 100 else None
    except Exception:
        pass
    return None


def _facebook_profile_id(url: str) -> str | None:
    """Extract Facebook profile/page ID from URL (e.g. facebook.com/Meta -> Meta, profile.php?id=123 -> 123)."""
    if not url or ("facebook.com" not in url.lower() and "fb.com" not in url.lower()):
        return None
    try:
        url = url if url.startswith("http") else "https://" + url
        p = urlparse(url)
        path = (p.path or "").strip().rstrip("/")
        query = (p.query or "").strip()
        if "profile.php" in path.lower() and "id=" in query:
            for part in query.split("&"):
                if part.strip().lower().startswith("id="):
                    return part.split("=", 1)[1].strip()[:100] or None
        if path:
            segment = path.strip("/").split("/")[0].split("?")[0].strip()
            if segment and segment.lower() not in ("profile.php", "pages", "watch", "groups", "events", "marketplace", "gaming", "login", "recover"):
                return segment[:100] if len(segment) < 100 else None
    except Exception:
        pass
    return None


def _instagram_username(url: str) -> str | None:
    """Extract Instagram username from URL (e.g. instagram.com/username -> username)."""
    if not url or "instagram.com" not in url.lower():
        return None
    try:
        url = url if url.startswith("http") else "https://" + url
        p = urlparse(url)
        path = (p.path or "").strip().rstrip("/")
        if path:
            segment = path.strip("/").split("/")[0].split("?")[0].strip()
            if segment and segment.lower() not in ("p", "reel", "reels", "stories", "explore", "accounts", "direct"):
                return segment[:80] if len(segment) < 80 else None
    except Exception:
        pass
    return None


def _twitter_username(url: str) -> str | None:
    """Extract Twitter/X username from URL (e.g. twitter.com/handle or x.com/handle -> handle)."""
    if not url:
        return None
    lower = url.lower()
    if "twitter.com" not in lower and "x.com" not in lower:
        return None
    try:
        url = url if url.startswith("http") else "https://" + url
        p = urlparse(url)
        path = (p.path or "").strip().rstrip("/")
        if path:
            segment = path.strip("/").split("/")[0].split("?")[0].strip()
            if segment and segment.lower() not in ("intent", "search", "home", "explore", "settings", "i", "share"):
                return segment[:80] if len(segment) < 80 else None
    except Exception:
        pass
    return None


def fetch_linkedin_profile_serpapi(profile_id: str, api_key: str) -> str | None:
    """Fetch public LinkedIn profile via SerpAPI; return formatted text or None."""
    if not profile_id or not api_key or len(profile_id) > 80:
        return None
    api_url = (
        "https://serpapi.com/search?engine=linkedin_profile"
        "&profile_id=" + quote(profile_id) + "&api_key=" + quote(api_key)
    )
    try:
        req = Request(api_url, headers={"User-Agent": "AskSiddhartha/1.0"})
        with urlopen(req, timeout=SERP_TIMEOUT) as resp:
            data = json.loads(resp.read().decode())
    except (URLError, HTTPError, OSError, json.JSONDecodeError):
        return None
    # Build readable summary from common SerpAPI LinkedIn profile fields
    parts = []
    if data.get("profile"):
        p = data.get("profile") or {}
        if p.get("name"):
            parts.append(f"Name: {p['name']}")
        if p.get("headline"):
            parts.append(f"Headline: {p['headline']}")
        if p.get("summary") or p.get("about"):
            parts.append(f"About: {p.get('summary') or p.get('about')}")
        if p.get("location"):
            parts.append(f"Location: {p['location']}")
        for key in ("experience", "education", "certifications", "skills"):
            val = p.get(key)
            if isinstance(val, list) and val:
                lines = []
                for i, item in enumerate(val[:15], 1):
                    if isinstance(item, dict):
                        line = " ".join(str(v) for k, v in item.items() if v and k != "link")
                        if line:
                            lines.append(f"  {i}. {line.strip()}")
                    elif isinstance(item, str):
                        lines.append(f"  {i}. {item}")
                if lines:
                    parts.append(f"{key.title()}:\n" + "\n".join(lines))
            elif isinstance(val, str) and val.strip():
                parts.append(f"{key.title()}: {val.strip()}")
    if not parts:
        # Fallback: use top-level profile-related keys or full response
        for key in ("profile", "experiences", "experience", "education", "certifications"):
            val = data.get(key)
            if val is not None:
                parts.append(f"{key}: {json.dumps(val, ensure_ascii=False)[:3000]}")
        if not parts:
            summary = json.dumps(data, ensure_ascii=False)[:8000]
            if summary and summary != "{}":
                parts.append(f"Profile data: {summary}")
    return "\n\n".join(parts) if parts else None


def fetch_facebook_profile_serpapi(profile_id: str, api_key: str) -> str | None:
    """Fetch public Facebook profile/page via SerpAPI; return formatted text or None."""
    if not profile_id or not api_key or len(profile_id) > 100:
        return None
    api_url = (
        "https://serpapi.com/search?engine=facebook_profile"
        "&profile_id=" + quote(profile_id) + "&api_key=" + quote(api_key)
    )
    try:
        req = Request(api_url, headers={"User-Agent": "AskSiddhartha/1.0"})
        with urlopen(req, timeout=SERP_TIMEOUT) as resp:
            data = json.loads(resp.read().decode())
    except (URLError, HTTPError, OSError, json.JSONDecodeError):
        return None
    pr = data.get("profile_results") or {}
    if not pr:
        return None
    parts = []
    if pr.get("name"):
        parts.append(f"Name: {pr['name']}")
    if pr.get("profile_type"):
        parts.append(f"Type: {pr['profile_type']}")
    if pr.get("profile_intro_text"):
        parts.append(f"Intro: {pr['profile_intro_text']}")
    if pr.get("category"):
        parts.append(f"Category: {pr['category']}")
    if pr.get("verified") is not None:
        parts.append(f"Verified: {pr['verified']}")
    if pr.get("followers"):
        parts.append(f"Followers: {pr['followers']}")
    if pr.get("likes"):
        parts.append(f"Likes: {pr['likes']}")
    if pr.get("address"):
        parts.append(f"Address: {pr['address']}")
    if pr.get("phone"):
        parts.append(f"Phone: {pr['phone']}")
    if pr.get("email"):
        parts.append(f"Email: {pr['email']}")
    if pr.get("business_hours"):
        parts.append(f"Hours: {pr['business_hours']}")
    if pr.get("url"):
        parts.append(f"URL: {pr['url']}")
    return "\n".join(parts) if parts else None


def fetch_social_profile_via_search(
    platform: str,
    username: str,
    api_key: str,
    want_latest_tweet: bool = False,
    want_latest_post: bool = False,
) -> str | None:
    """Use Google search to get public info about an Instagram or Twitter/X profile (no dedicated SerpAPI engine)."""
    if not username or not api_key or platform not in ("instagram", "twitter"):
        return None
    recency = "m" if (want_latest_tweet or want_latest_post) else None
    if platform == "instagram":
        query = f"{username} latest post instagram" if want_latest_post else f"{username} {platform} profile"
    elif want_latest_tweet:
        query = f"{username} latest tweet"
    else:
        query = f"{username} Twitter X profile"
    result = fetch_web_search(query, api_key, recency_filter=recency) if recency else fetch_web_search(query, api_key)
    if not result and recency:
        result = fetch_web_search(query, api_key)
    return result


# Real-time web search via SerpAPI (Google SERP)
SERP_MAX_RESULTS = 8
SERP_TIMEOUT = 10

# Phrases that indicate user wants the last/most recent post by that account (may be from long ago)
LATEST_POST_PHRASES = (
    "latest post", "latest tweet", "recent tweet", "recent post", "last post", "last tweet",
    "newest tweet", "newest post", "most recent post", "most recent tweet",
)


def _extract_name_and_platform_for_latest(message: str) -> tuple[str | None, str | None]:
    """Extract (name, platform) when user asks for latest post. Platform: twitter, instagram, facebook, linkedin."""
    lower = (message or "").strip().lower()
    if not any(p in lower for p in LATEST_POST_PHRASES):
        return None, None
    platform = None
    if "twitter" in lower or "tweet" in lower or "x.com" in lower:
        platform = "twitter"
    elif "instagram" in lower or "insta" in lower:
        platform = "instagram"
    elif "facebook" in lower or "fb " in lower or "fb." in lower:
        platform = "facebook"
    elif "linkedin" in lower or "linked in" in lower:
        platform = "linkedin"
    if not platform:
        return None, None
    name = lower
    for prefix in (
        "show the latest post by ", "show latest post by ", "get the latest post by ",
        "latest post by ", "recent post by ", "last post by ", "last tweet by ",
        "newest post by ", "latest tweet by ", "recent tweet by ",
        "show the last post by ", "last post by ", "get the last post by ",
    ):
        if name.startswith(prefix):
            name = name[len(prefix):].strip()
            break
    for suffix in (
        " on twitter", " on x", " from twitter", " from x", " twitter", "'s twitter", " twitter account",
        " on instagram", " from instagram", " instagram", "'s instagram", " instagram account",
        " on facebook", " from facebook", " facebook", "'s facebook", " facebook page",
        " on linkedin", " from linkedin", " linkedin", "'s linkedin", " linkedin profile",
    ):
        if name.endswith(suffix):
            name = name[: -len(suffix)].strip()
    name = re.sub(r"\s+", " ", name).strip()
    if not name or len(name) < 2:
        return None, None
    return name, platform


def _refine_query_for_latest_tweet(message: str) -> str | None:
    """If the user asks for latest post on Twitter/Instagram/Facebook/LinkedIn, return a search query for that."""
    name, platform = _extract_name_and_platform_for_latest(message)
    if not name or not platform:
        return None
    if platform == "twitter":
        return f"{name} latest tweet"
    return f"{name} latest post {platform}"


def fetch_web_search(query: str, api_key: str, recency_filter: str | None = None) -> str | None:
    """Call SerpAPI for Google search; return formatted snippet string or None.
    recency_filter: 'w' = past week, 'm' = past month; adds tbs=qdr:w/m to bias toward recent results."""
    if not query or not api_key or len(query.strip()) < 2:
        return None
    query = query.strip()[:300]
    url = "https://serpapi.com/search?q=" + quote(query) + "&api_key=" + quote(api_key) + "&num=" + str(SERP_MAX_RESULTS)
    if recency_filter in ("w", "m"):
        url += "&tbs=qdr:" + recency_filter
    try:
        req = Request(url, headers={"User-Agent": "AskSiddhartha/1.0"})
        with urlopen(req, timeout=SERP_TIMEOUT) as resp:
            data = json.loads(resp.read().decode())
        results = data.get("organic_results") or []
        if not results:
            return None
        parts = []
        for i, r in enumerate(results[:SERP_MAX_RESULTS], 1):
            title = (r.get("title") or "").strip()
            link = (r.get("link") or "").strip()
            snippet = (r.get("snippet") or "").strip()
            date_str = (r.get("date") or "").strip()
            if title or snippet:
                line = f"{i}. {title}\n   {snippet}\n   URL: {link}"
                if date_str:
                    line += f"\n   Date: {date_str}"
                parts.append(line)
        return "\n\n".join(parts) if parts else None
    except (URLError, HTTPError, OSError, json.JSONDecodeError, KeyError):
        return None


def get_openai_client():
    try:
        from openai import OpenAI
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            return None
        return OpenAI(api_key=api_key)
    except ImportError:
        return None


def fetch_person_image(person_name: str) -> str | None:
    """Fetch a real photo of a person from Wikipedia; return base64 string or None."""
    if not person_name or len(person_name) > 200:
        return None
    try:
        search_url = (
            "https://en.wikipedia.org/w/api.php"
            "?action=query&list=search&srsearch=" + quote(person_name) +
            "&format=json&srlimit=1"
        )
        req = Request(search_url, headers={"User-Agent": "AskSiddhartha/1.0 (https://github.com/)"})
        with urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read().decode())
        pages = data.get("query", {}).get("search", [])
        if not pages:
            return None
        title = pages[0].get("title", "")
        if not title:
            return None
        img_url_api = (
            "https://en.wikipedia.org/w/api.php"
            "?action=query&titles=" + quote(title) +
            "&prop=pageimages&format=json&pithumbsize=800"
        )
        req2 = Request(img_url_api, headers={"User-Agent": "AskSiddhartha/1.0"})
        with urlopen(req2, timeout=8) as resp2:
            data2 = json.loads(resp2.read().decode())
        query = data2.get("query", {}).get("pages", {})
        page = next((p for p in query.values() if "pageimage" in p), None)
        if not page:
            return None
        thumb = page.get("thumbnail")
        if not thumb:
            return None
        image_url = thumb.get("source")
        if not image_url:
            return None
        req3 = Request(image_url, headers={"User-Agent": "AskSiddhartha/1.0"})
        with urlopen(req3, timeout=10) as resp3:
            img_bytes = resp3.read()
        return base64.standard_b64encode(img_bytes).decode("ascii")
    except (URLError, HTTPError, json.JSONDecodeError, KeyError, OSError):
        return None


def search_person_image_web(person_name: str) -> str | None:
    """Search the web for a person's photo when Wikipedia has none; return base64 or None."""
    if not person_name or len(person_name) > 200:
        return None
    try:
        from duckduckgo_search import DDGS
        # For single person, use a more specific query to get correct/official photos
        is_together = " together" in person_name.lower() or " and " in person_name.lower()
        search_query = person_name if is_together else (person_name + " official photo")
        with DDGS() as ddgs:
            results = list(ddgs.images(keywords=search_query, type_image="photo", max_results=12))
        if not results:
            return None
        # Prefer Wikipedia/Wikimedia URLs so we get the correct person's image
        def rank_url(item: dict) -> tuple:
            url = (item.get("image") or item.get("url") or "").lower()
            if "wikipedia" in url or "wikimedia" in url:
                return (0, url)
            return (1, url)
        results = sorted(results, key=rank_url)
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
        max_bytes = 4 * 1024 * 1024  # 4 MB cap
        for item in results:
            image_url = (item.get("image") or item.get("url") or "").strip()
            if not image_url or not image_url.startswith("http"):
                continue
            try:
                req = Request(image_url, headers=headers)
                with urlopen(req, timeout=12) as resp:
                    raw = resp.read(max_bytes)
                if len(raw) < 500:
                    continue
                return base64.standard_b64encode(raw).decode("ascii")
            except (URLError, HTTPError, OSError):
                continue
        return None
    except Exception:
        return None


def generate_image(prompt: str) -> str | None:
    """Generate an image via DALL-E; return base64 string or None."""
    client = get_openai_client()
    if not client:
        return None
    try:
        r = client.images.generate(
            model=os.environ.get("OPENAI_IMAGE_MODEL", "dall-e-3"),
            prompt=prompt[:4000],
            n=1,
            size="1024x1024",
            response_format="b64_json",
            quality="standard",
        )
        if r.data and len(r.data) > 0:
            return r.data[0].b64_json
    except Exception:
        pass
    return None


def chat_with_llm(messages: list[dict], lang_name: str, session_id: str, is_voice_input: bool = False, link_context: str | None = None) -> str | dict:
    """Call LLM with session messages and any uploaded files. Returns assistant reply."""
    client = get_openai_client()
    if not client:
        return (
            "Hello! I'm your multilingual assistant. "
            "Set OPENAI_API_KEY in your environment to enable AI replies. "
            f"You're using: {lang_name}. Session memory is active for this chat only."
        )

    today = datetime.now(timezone.utc).strftime("%d %B %Y")
    ist = timezone(timedelta(hours=5, minutes=30))
    now_ist = datetime.now(ist)
    ist_date_str = now_ist.strftime("%d %B %Y")
    ist_time_str = now_ist.strftime("%I:%M %p IST")
    reply_lang = (lang_name or "English").strip()
    system = (
        f"You are a helpful, friendly assistant. You must write every response only in {reply_lang}. "
        "The user may type in any language (e.g. Hindi, Tamil, English); you must still reply only in the selected language. "
        "Uploaded documents or images may be in any language; give your answer in the user's selected language only. "
        "When the user shares a link, you may receive extracted text from that page. LinkedIn and Facebook profile URLs may be resolved to profile data; Instagram and Twitter/X URLs use web search results. Use any provided profile or search data to answer questions about the person or account. If you receive '[URL: ...]' with page or profile content below it, you MUST use that content to answer—do NOT say you cannot access external links. If you receive a '[Note: ... could not retrieve the page content]' then say you could not load the page and suggest pasting details or trying again. Reply in the selected language. "
        "When you receive '[Web search results for your query]' or '[Web search results that may mention the site or your query]' below, use them to give current, real-time answers; summarize or cite from the results as needed. "
        "When you see '[IMPORTANT: The user asked for the LATEST or LAST post/tweet...]', you MUST use ONLY Result 1 (the first result) as the latest post unless it clearly is not a post/tweet. If results include 'Date:', pick the result with the most recent date. Do NOT use result 2 or 3 as the 'latest'. Present the post text and date clearly. 'Latest' can be from long ago; still show it as their latest. "
        "Do not say things like 'Paste a link and type your question, or just ask anything. I'll read the page and answer.' or similar link-prompt lines. "
        "Give helpful, detailed answers: use a few sentences or a short paragraph when useful. Do not limit yourself to one or two lines unless the question is trivial. "
        "Do not repeat the date, previous answers, or facts unless the user clearly asks again. "
        "Do not say 'I apologize', 'It appears you may have made a mistake', or similar unless the user is correcting a real error. "
        "For very short or unclear messages (e.g. '26.', 'ok'), reply in one short line or ask how you can help—do not re-explain or apologize. "
        "If the user asks for a photo or picture of a specific real person (actor, singer, politician, sportsperson, etc.), "
        "reply with exactly one line: [SEARCH_PERSON_IMAGE: full official name, role] e.g. [SEARCH_PERSON_IMAGE: Shah Rukh Khan, actor] or [SEARCH_PERSON_IMAGE: Narendra Modi, politician]. "
        "Use the same role the user asked for (actor, singer, politician, etc.)—do not substitute a different role. "
        "If the user asks for a photo of two or more people together, use [SEARCH_PERSON_IMAGE: Person A and Person B together]. "
        "Then on the next line write a short caption in the user's language. In the caption, use the same role the user asked for (e.g. if they asked for an actor, write 'actor', not 'singer' or something else). Use full official names. "
        "If the user asks for a photo or image of something that is NOT a specific real person (e.g. sunset, a cat, a generic scene), "
        "reply with exactly one line: [GENERATE_IMAGE: detailed description in English for image generation]. Then a short caption on the next line. "
        "Do not use [SEARCH_PERSON_IMAGE] for fictional characters or generic descriptions; use [GENERATE_IMAGE] for those. "
        "For normal questions do not use these tags. "
        "When the user asks for code (any programming language): write the code itself always in English (keywords, syntax, variable/function names as in standard practice). "
        "Explanations before or after the code, and comments inside the code, may be in the user's selected language ({lang_name}). "
        f"Today's date is {today}. The current date and time in India (IST) right now is: {ist_date_str}, {ist_time_str}. When the user asks for the current date, time, or date and time in India, use this exact value—do not use search results or guess. This conversation has no memory outside this session."
    )
    api_messages = [{"role": "system", "content": system}]
    if is_voice_input and lang_name and lang_name.lower() != "english":
        api_messages.append({
            "role": "system",
            "content": f"The user's message is from voice input (often transcribed in English). You must reply only in {lang_name}.",
        })

    uploads = SESSION_UPLOADS.get(session_id, [])
    if uploads:
        content_parts = []
        doc_texts = []
        for u in uploads:
            if u.get("type") == "document" and u.get("content"):
                doc_texts.append(f"[{u.get('filename', 'document')}]\n{u['content']}")
            elif u.get("type") == "image" and u.get("content"):
                content_parts.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{u.get('mime', 'image/png')};base64,{u['content']}"},
                })
        if doc_texts:
            content_parts.insert(0, {
                "type": "text",
                "text": "The user has shared these documents (they may be in any language). Use this content to answer. Reply only in the user's selected language.\n\n" + "\n\n---\n\n".join(doc_texts),
            })
        elif content_parts and all(p.get("type") == "image_url" for p in content_parts):
            content_parts.insert(0, {"type": "text", "text": "The user has shared these images (they may be in any language). Use them to answer. Reply only in the user's selected language."})
        if content_parts:
            if len(content_parts) == 1 and content_parts[0].get("type") == "text":
                api_messages.append({"role": "user", "content": content_parts[0]["text"]})
            else:
                api_messages.append({"role": "user", "content": content_parts})

    web_search_context = None
    if messages and messages[-1].get("role") == "user":
        last_query = (messages[-1].get("content") or "").strip()[:300]
        serp_key = os.environ.get("SERPAPI_API_KEY") or os.environ.get("GOOGLE_SERP_API_KEY")
        if serp_key and last_query:
            # For "latest post/tweet" on Twitter/Instagram/etc., search with recency bias and strict instructions
            refined = _refine_query_for_latest_tweet(last_query)
            search_query = refined if refined else last_query
            if refined:
                web_search_context = fetch_web_search(search_query, serp_key, recency_filter="m")
                if not web_search_context:
                    web_search_context = fetch_web_search(search_query, serp_key)
            else:
                web_search_context = fetch_web_search(search_query, serp_key)
            if web_search_context and refined:
                web_search_context = (
                    "[IMPORTANT: The user asked for the LATEST or LAST post/tweet from this account. "
                    "You MUST use ONLY the first result (Result 1) as the latest post—unless it clearly has no post/tweet content. "
                    "If results include a 'Date:' line, the result with the most recent date is the latest; prefer that one. "
                    "Do NOT use result 2, 3, etc. as the 'latest' unless result 1 is not a post. "
                    "Present the post text and date clearly. 'Latest' can be from long ago; still show it as their latest.]\n\n"
                    + web_search_context
                )

    for i, m in enumerate(messages):
        content = m["content"]
        is_last_user = i == len(messages) - 1 and m.get("role") == "user"
        if is_last_user and reply_lang.lower() == "english":
            content = "[Instruction: You must respond in English only. Do not use Hindi or any other language.]\n\n" + content
        if (web_search_context or link_context) and is_last_user:
            prepend = ""
            if web_search_context:
                prepend += "[Web search results for your query:]\n\n" + web_search_context + "\n\n"
            if link_context:
                intro = (
                    "[The user shared a link; we could not load the page. See note below.]\n\n"
                    if link_context.strip().startswith("[Note:")
                    else "[Content extracted from the shared link(s). Use this to answer the user's question.]\n\n"
                )
                prepend += intro + link_context + "\n\n"
            content = prepend + "---\n\nUser's question: " + content
        api_messages.append({"role": m["role"], "content": content})

    try:
        max_completion_tokens = 900 if (link_context or web_search_context) else 600
        r = client.chat.completions.create(
            model=os.environ.get("OPENAI_MODEL", "gpt-5.3-chat-latest"),
            messages=api_messages,
            max_completion_tokens=max_completion_tokens,
        )
        text = (r.choices[0].message.content or "").strip()
        # When English is selected but the model replied in Indic script, retry once with a strict English-only instruction
        if reply_lang.lower() == "english" and _reply_contains_indic(text):
            api_messages.append({"role": "assistant", "content": text})
            api_messages.append({"role": "user", "content": "You must respond in English only. Give the same answer again in English."})
            try:
                r2 = client.chat.completions.create(
                    model=os.environ.get("OPENAI_MODEL", "gpt-5.3-chat-latest"),
                    messages=api_messages,
                    max_completion_tokens=max_completion_tokens,
                )
                text = (r2.choices[0].message.content or "").strip()
            except Exception:
                pass
        if "[SEARCH_PERSON_IMAGE:" in text:
            match = re.search(r"\[SEARCH_PERSON_IMAGE:\s*([^\]]+)\]", text)
            if match:
                raw = match.group(1).strip()
                # Parse "Name, role" or "Name" or "Person A and Person B together"
                if ", " in raw and " together" not in raw and " and " not in raw.split(", ", 1)[0]:
                    person_name, role = raw.split(", ", 1)
                    person_name = person_name.strip()
                    role = role.strip()
                else:
                    person_name = raw
                    role = ""
                caption = re.sub(r"\[SEARCH_PERSON_IMAGE:[^\]]+\]\s*", "", text).strip()
                if "\n" in caption:
                    caption = caption.split("\n")[-1].strip()
                if not caption:
                    caption = f"{person_name} ({role})" if role else "Here's the photo."
                elif role and role.lower() not in caption.lower():
                    caption = f"{person_name} ({role})"
                search_query = f"{person_name} {role}".strip() if role else person_name
                is_together_query = " together" in person_name.lower() or " and " in person_name.lower()
                b64 = None if is_together_query else fetch_person_image(person_name)
                if not b64:
                    b64 = search_person_image_web(search_query)
                if b64:
                    return {"reply": caption, "image_b64": b64}
                return {"reply": caption + " (Photo not found.)", "image_b64": None}
        if "[GENERATE_IMAGE:" in text:
            match = re.search(r"\[GENERATE_IMAGE:\s*([^\]]+)\]", text)
            if match:
                prompt = match.group(1).strip()
                caption = re.sub(r"\[GENERATE_IMAGE:[^\]]+\]\s*", "", text).strip()
                if "\n" in caption:
                    caption = caption.split("\n")[-1].strip()
                if not caption:
                    caption = "Here’s the image."
                b64 = generate_image(prompt)
                if b64:
                    return {"reply": caption or "Here’s the image.", "image_b64": b64}
                return {"reply": caption or "I couldn’t generate the image. Please try again.", "image_b64": None}
        return text
    except Exception as e:
        return f"Sorry, I couldn't get a reply right now: {str(e)}"


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/session", methods=["GET"])
def api_session():
    sid = get_or_create_session_id()
    return jsonify({"session_id": sid, "messages": get_messages(sid)})


@app.route("/api/transliterate", methods=["POST"])
def api_transliterate():
    data = request.get_json() or {}
    text = (data.get("text") or "").strip()
    lang_code = (data.get("lang_code") or "hi").lower()
    if not text or lang_code == "en":
        return jsonify({"original": text, "transliterated": "", "direction": ""})
    if not HAS_TRANSLITERATION or lang_code not in SCRIPT_MAP:
        return jsonify({"original": text, "transliterated": "", "direction": ""})
    # If user typed in Indic script, show Roman; else show Indic
    if _has_indic_script(text, lang_code):
        out = transliterate_to_roman(text, lang_code)
        return jsonify({"original": text, "transliterated": out, "direction": "to_roman"})
    out = transliterate_to_indic(text, lang_code)
    return jsonify({"original": text, "transliterated": out, "direction": "to_indic"})


@app.route("/api/chat", methods=["POST"])
def api_chat():
    data = request.get_json() or {}
    user_message = (data.get("message") or "").strip()
    lang_code = (data.get("lang_code") or "en").lower()
    lang_name = (data.get("lang_name") or "English")
    is_voice_input = data.get("is_voice_input", False)

    if not user_message:
        return jsonify({"error": "Message is required"}), 400

    link_context = None
    shared_url = (data.get("shared_url") or "").strip()
    urls = extract_urls_from_text(user_message)
    if shared_url and shared_url not in urls:
        urls.insert(0, shared_url)
    if urls:
        urls = urls[:MAX_LINKS_PER_MESSAGE]
        parts = []
        serp_key = os.environ.get("SERPAPI_API_KEY") or os.environ.get("GOOGLE_SERP_API_KEY")
        want_latest = any(
                p in (user_message or "").lower()
                for p in LATEST_POST_PHRASES
            )
        for url in urls:
            text = None
            if serp_key:
                lid = _linkedin_profile_id(url)
                if lid:
                    if want_latest:
                        text = fetch_web_search(f"{lid} latest post linkedin", serp_key, recency_filter="m")
                        if not text:
                            text = fetch_web_search(f"{lid} latest post linkedin", serp_key)
                    if not text:
                        text = fetch_linkedin_profile_serpapi(lid, serp_key)
                if not text:
                    fid = _facebook_profile_id(url)
                    if fid:
                        if want_latest:
                            text = fetch_web_search(f"{fid} latest post facebook", serp_key, recency_filter="m")
                            if not text:
                                text = fetch_web_search(f"{fid} latest post facebook", serp_key)
                        if not text:
                            text = fetch_facebook_profile_serpapi(fid, serp_key)
                if not text:
                    iid = _instagram_username(url)
                    if iid:
                        text = fetch_social_profile_via_search(
                            "instagram", iid, serp_key, want_latest_post=want_latest
                        )
                if not text:
                    tid = _twitter_username(url)
                    if tid:
                        text = fetch_social_profile_via_search(
                            "twitter", tid, serp_key, want_latest_tweet=want_latest
                        )
            if not text:
                text = fetch_link_content(url)
            if text:
                parts.append(f"[URL: {url}]\n{text}")
        if parts:
            link_context = "\n\n".join(parts)
        else:
            # Fetch failed for all URLs (travel/booking sites often use JavaScript and block bots)
            link_context = (
                "[Note: The user shared these link(s) but we could not retrieve the page content. "
                "Many travel and booking sites (e.g. MakeMyTrip, booking engines) load content via JavaScript "
                "or block automated access, so we cannot read them directly. "
                "Do NOT say you cannot access external links. Use any web search results below if provided; "
                "otherwise say you could not load the page and suggest they ask for general estimates or paste details. "
                "Link(s): " + ", ".join(urls) + "]"
            )
            # When we couldn't load the page, run a targeted web search (user query + site domain) so the bot has something to work with
            serp_key = os.environ.get("SERPAPI_API_KEY") or os.environ.get("GOOGLE_SERP_API_KEY")
            if serp_key and user_message:
                domains = list({_domain_from_url(u) for u in urls if _domain_from_url(u)})
                query_text = user_message.strip()
                for u in urls:
                    query_text = query_text.replace(u, " ").strip()
                query_text = " ".join(query_text.split())[:200]
                search_query = (query_text + (" " + " ".join(domains)) if domains else query_text).strip()
                if search_query and len(search_query) >= 3:
                    site_search = fetch_web_search(search_query, serp_key)
                    if site_search:
                        link_context += "\n\n[Web search results that may mention the site or your query:]\n\n" + site_search

    sid = get_or_create_session_id()
    add_message(sid, "user", user_message)
    messages = get_messages(sid)
    result = chat_with_llm(messages, lang_name, sid, is_voice_input=is_voice_input, link_context=link_context)
    if isinstance(result, dict):
        reply = result.get("reply", "")
        image_b64 = result.get("image_b64")
        add_message(sid, "assistant", reply)
        return jsonify({
            "reply": reply,
            "image_b64": image_b64,
            "is_voice_input": is_voice_input,
            "session_id": sid,
        })
    add_message(sid, "assistant", result)
    return jsonify({
        "reply": result,
        "is_voice_input": is_voice_input,
        "session_id": sid,
    })


@app.route("/api/clear", methods=["POST"])
def api_clear():
    """Clear session memory and uploads; start a new session so memory is gone (used for New chat and language change)."""
    sid = get_or_create_session_id()
    SESSION_MEMORY[sid] = []
    SESSION_UPLOADS[sid] = []
    session.pop("session_id", None)  # Next request gets a new session_id → fresh conversation
    return jsonify({"ok": True})


ALLOWED_UPLOAD_EXTENSIONS = {"pdf", "doc", "docx", "png", "jpg", "jpeg", "gif", "webp"}
MAX_UPLOAD_MB = 10


@app.route("/api/upload", methods=["POST"])
def api_upload():
    """Accept screenshot (image) or PDF/Word document; store per session for chat context."""
    sid = get_or_create_session_id()
    if sid not in SESSION_UPLOADS:
        SESSION_UPLOADS[sid] = []

    if "file" not in request.files and "files" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    files = request.files.getlist("files") or request.files.getlist("file") or ([request.files["file"]] if request.files.get("file") else [])
    added = []
    for f in files:
        if not f or not f.filename:
            continue
        ext = (f.filename.rsplit(".", 1)[-1] or "").lower()
        if ext not in ALLOWED_UPLOAD_EXTENSIONS:
            continue
        data = f.read()
        if len(data) > MAX_UPLOAD_MB * 1024 * 1024:
            continue
        filename = f.filename
        if ext in ("pdf", "doc", "docx"):
            if ext == "pdf":
                text = extract_pdf_text(data)
            else:
                text = extract_docx_text(data) if ext == "docx" else ""
            if not text and ext == "doc":
                text = "(Could not extract text from .doc; try saving as .docx)"
            SESSION_UPLOADS[sid].append({"type": "document", "content": text[:50000], "filename": filename})
            added.append({"filename": filename, "type": "document"})
        else:
            b64 = base64.standard_b64encode(data).decode("ascii")
            mime = "image/png" if ext == "png" else "image/jpeg" if ext in ("jpg", "jpeg") else "image/gif" if ext == "gif" else "image/webp"
            SESSION_UPLOADS[sid].append({"type": "image", "content": b64, "filename": filename, "mime": mime})
            added.append({"filename": filename, "type": "image"})

    return jsonify({"ok": True, "files": added})


@app.route("/api/uploads", methods=["GET"])
def api_uploads_list():
    """Return list of uploaded filenames for current session."""
    sid = get_or_create_session_id()
    files = [{"filename": u.get("filename", ""), "type": u.get("type", "")} for u in SESSION_UPLOADS.get(sid, [])]
    return jsonify({"files": files})


@app.route("/api/uploads/clear", methods=["POST"])
def api_uploads_clear():
    """Clear uploaded files for current session."""
    sid = get_or_create_session_id()
    SESSION_UPLOADS[sid] = []
    return jsonify({"ok": True})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=os.environ.get("FLASK_DEBUG", "0") == "1")
